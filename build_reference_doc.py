from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page margins ---
section = doc.sections[0]
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val.get('val', 'single'))
            el.set(qn('w:sz'),    str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', 'auto'))
            borders.append(el)
    tcPr.append(borders)

def add_run_with_color(para, text, bold=False, font_size=11,
                        color=None, font_name='Calibri'):
    run = para.add_run(text)
    run.bold      = bold
    run.font.size = Pt(font_size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run

# ── Intro callout box (single-cell table) ──────────────────────────────────

intro_tbl = doc.add_table(rows=1, cols=1)
intro_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
intro_tbl.style     = 'Table Grid'

cell = intro_tbl.cell(0, 0)
set_cell_bg(cell, 'E8F0F7')

# thick left border accent
tc   = cell._tc
tcPr = tc.get_or_add_tcPr()
borders = OxmlElement('w:tcBorders')
for side in ('top', 'bottom', 'right'):
    el = OxmlElement(f'w:{side}')
    el.set(qn('w:val'),   'single')
    el.set(qn('w:sz'),    '4')
    el.set(qn('w:space'), '0')
    el.set(qn('w:color'), 'BDD0E0')
    borders.append(el)
left_el = OxmlElement('w:left')
left_el.set(qn('w:val'),   'single')
left_el.set(qn('w:sz'),    '18')
left_el.set(qn('w:space'), '0')
left_el.set(qn('w:color'), '2A6496')
borders.append(left_el)
tcPr.append(borders)

cell.paragraphs[0].clear()
p = cell.paragraphs[0]
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
p.paragraph_format.left_indent  = Inches(0.1)

add_run_with_color(p, 'The adaptation projects identified in this plan are informed by the ', font_size=11)
add_run_with_color(p, '2026 Vulnerability Assessment and Stormwater Model Update Technical Report',
                   bold=True, font_size=11)
add_run_with_color(p, (', and are aligned with projects, recommendations, and priorities identified '
                       'in eleven existing plans and studies for Village of Estero. For ease of '
                       'reference, each source has been assigned a number below and is cited by '
                       'that number in the project sheets.'), font_size=11)

doc.add_paragraph()  # spacer

# ── Source reference table ─────────────────────────────────────────────────

NAVY      = '1A4F7A'
WHITE     = 'FFFFFF'

sources = [
    (1,  '2018', 'Village of Estero Comprehensive Plan',                                             '2A6496', 'EAF3FB'),
    (2,  '2018', 'Village of Estero Stormwater Master Plan',                                         '2A6496', 'DDEEF8'),
    (3,  '2023', 'Village of Estero Flood Mitigation Project Memorandum',                            '27AE60', 'E6F4EA'),
    (4,  '2025', 'Village of Estero Hazard Mitigation Grant Program Memorandum: Pond Design',        'D4840D', 'FEF9E7'),
    (5,  '2022', 'Lee County Joint Local Mitigation Strategy',                                       '8E44AD', 'FDF2F8'),
    (6,  '2024', 'Village of Estero Narrative Provided to FEMA',                                     'C0392B', 'FEF9E7'),
    (7,  '2018', 'Estero River, Imperial River, & Springs Creek Storm Assessment Summary Report',    '2A6496', 'EAF3FB'),
    (8,  '2021', 'South Lee County Watershed Initiative Hydrological Modeling Project Final Report', '1A8A7A', 'E6F4EA'),
    (9,  '2025', 'Lee County Vulnerability Assessment Report',                                       'D4840D', 'FEF9E7'),
    (10, '2025', 'Lee County Coastal Flood Adaptation Plan',                                         'D4840D', 'FDF2F8'),
    (11, '2025', 'Village of Estero Progress Report on Local Mitigation Strategy',                   'D4840D', 'E6F4EA'),
]

tbl = doc.add_table(rows=1 + len(sources), cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style     = 'Table Grid'

# column widths: #, Year tag, Document title
col_widths = [Inches(0.55), Inches(1.05), Inches(5.4)]
for row in tbl.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_widths[i]

# header row
hdr_cells = tbl.rows[0].cells
hdr_data   = ['#', 'YEAR', 'SOURCE DOCUMENT']
for i, (cell, txt) in enumerate(zip(hdr_cells, hdr_data)):
    set_cell_bg(cell, NAVY)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    add_run_with_color(p, txt, bold=True, font_size=10.5, color=WHITE)

# data rows
for row_idx, (num, year, title, badge_color, row_bg) in enumerate(sources):
    row_cells = tbl.rows[row_idx + 1].cells

    # col 0 — number badge
    num_cell = row_cells[0]
    set_cell_bg(num_cell, badge_color)
    num_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p0 = num_cell.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(4)
    p0.paragraph_format.space_after  = Pt(4)
    add_run_with_color(p0, str(num), bold=True, font_size=13, color=WHITE)

    # col 1 — year pill
    year_cell = row_cells[1]
    set_cell_bg(year_cell, badge_color)
    year_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p1 = year_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after  = Pt(4)
    add_run_with_color(p1, year, bold=True, font_size=10, color=WHITE)

    # col 2 — document title
    title_cell = row_cells[2]
    set_cell_bg(title_cell, row_bg)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p2 = title_cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after  = Pt(4)
    p2.paragraph_format.left_indent  = Inches(0.05)
    add_run_with_color(p2, title, font_size=11, color='1A2A3A')

out = '/home/user/HERMES/estero_source_references.docx'
doc.save(out)
print(f'Saved: {out}')
