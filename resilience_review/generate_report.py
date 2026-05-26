from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Style helpers ─────────────────────────────────────────────────────────────
def font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold  = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    if level == 1:
        font(run, size=13, bold=True, color=(31, 73, 125))
    elif level == 2:
        font(run, size=11, bold=True, color=(55, 96, 146))
    else:
        font(run, size=11, bold=True, italic=True)

def body(text, space_after=6, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    font(run, italic=italic, color=color)
    return p

def note(text):
    """Italic gray note — used for 'this needs GIS / not yet verified' callouts."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_after  = Pt(5)
    run = p.add_run("Note: " + text)
    font(run, italic=True, color=(89, 89, 89))

def quote(text, source):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f'"{text}"')
    font(r1, italic=True, color=(64, 64, 64))
    r2 = p.add_run(f"\n— {source}")
    font(r2, size=9.5, italic=True, color=(127, 127, 127))

def bullet(text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + indent * 0.2)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    font(run)

def mixed(label, rest, indent=False):
    """Bullet with bold label + normal continuation."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.45)
    r1 = p.add_run(label)
    font(r1, bold=True)
    r2 = p.add_run(rest)
    font(r2)

def shaded_row(row, hex_fill="1F497D"):
    for cell in row.cells:
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_fill)
        tcPr.append(shd)

def table_header(tbl, cols):
    row = tbl.add_row()
    for i, text in enumerate(cols):
        cell = row.cells[i]
        cell.text = ""
        p   = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        font(run, size=9.5, bold=True, color=(255, 255, 255))
    shaded_row(row)
    return row

def table_row(tbl, vals, bold=False):
    row = tbl.add_row()
    for i, text in enumerate(vals):
        cell = row.cells[i]
        cell.text = ""
        p   = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        font(run, size=9.5, bold=bold)
    return row

# ═════════════════════════════════════════════════════════════════════════════
# COVER
# ═════════════════════════════════════════════════════════════════════════════
def centered(text, size, bold=False, italic=False, color=(0,0,0), space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    font(run, size=size, bold=bold, italic=italic, color=color)

centered("Northeast Florida Military Installation Resilience Review",
         size=16, bold=True, color=(31,73,125), space_before=20, space_after=4)
centered("Energy Infrastructure — Preliminary Research Findings and Investigation Routes",
         size=13, bold=True, color=(55,96,146), space_after=4)
centered(
    "Naval Air Station Jacksonville  |  Naval Station Mayport\n"
    "MCSF Blount Island  |  Camp Blanding JTF\n"
    "Duval County and Clay County, Florida",
    size=10, color=(89,89,89), space_after=4)
centered("Prepared for Supervisor Review — May 2026", size=10, italic=True,
         color=(89,89,89), space_after=12)

# ── Scope notice ──────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.left_indent  = Inches(0.3)
p.paragraph_format.right_indent = Inches(0.3)
p.paragraph_format.space_after  = Pt(10)
shd = OxmlElement("w:shd")
shd.set(qn("w:val"),   "clear")
shd.set(qn("w:color"), "auto")
shd.set(qn("w:fill"),  "EAF0F8")
p._p.get_or_add_pPr().append(shd)
run = p.add_run(
    "Scope and Honesty Notice\n"
    "This document reflects where the intern research team currently stands. "
    "Vulnerabilities cited are drawn from confirmed MIRR project findings — the vulnerability "
    "assessment, stakeholder documents, and the APTIM evaluation memo — not from independent "
    "GIS analysis conducted by this team. GIS analysis has not yet been performed. "
    "Where that matters — where a claim is ours rather than the MIRR's — this document says so "
    "explicitly. Routes for further investigation are raised as questions, not recommendations."
)
font(run, size=10, italic=True, color=(31,73,125))
run.bold = False
run2 = p.runs[0]  # re-bold the first line
p.clear()
r1 = p.add_run("Scope and Honesty Notice\n")
font(r1, size=10, bold=True, color=(31,73,125))
r2 = p.add_run(
    "This document reflects where the intern research team currently stands. "
    "Vulnerabilities cited are drawn from confirmed MIRR project findings — the vulnerability "
    "assessment, stakeholder documents, and the APTIM evaluation memo — not from independent "
    "GIS analysis conducted by this team. GIS analysis has not yet been performed. "
    "Where that matters — where a claim is ours rather than the MIRR's — this document says so "
    "explicitly. Routes for further investigation are raised as questions, not recommendations."
)
font(r2, size=10, italic=True, color=(31,73,125))

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHAT THE MIRR HAS CONFIRMED
# ═════════════════════════════════════════════════════════════════════════════
heading("1. What the MIRR Has Already Confirmed", level=1)
body(
    "The following findings come directly from MIRR vulnerability assessment documents and "
    "are cited here as the established baseline — not as intern-generated conclusions."
)

heading("1.1  The Bottom-Line Finding", level=2)
quote(
    "At the asset systems level there is no immediate risk to mission other than single points "
    "of failure at each installation.",
    "MIRR Vulnerability Assessment, January 2026 Steering Committee"
)
body(
    "All four installations rely on a single electrical substation with no backup supply path. "
    "Utility IGSAs at NAS Jacksonville and NS Mayport cover operations and maintenance only — "
    "not resilience or redundancy. No formal resilience coordination exists between Camp Blanding "
    "and FPL."
    "  (MIRR Mutual Support Assessment, February 2026 TAC)"
)

heading("1.2  Confirmed Energy Risk Ratings", level=2)
body("From the January 2026 Steering Committee vulnerability assessment table:")

tbl = doc.add_table(rows=1, cols=5)
tbl.style = "Table Grid"
table_header(tbl, ["Installation", "Substation / Owner", "Risk to Mission",
                   "Urgency", "Confirmed Hazard Exposure"])
table_row(tbl, ["NS Mayport",      "JEA Mayport Sub (south of base)", "HIGH",   "Immediate",  "Storm surge, rainfall flooding"])
table_row(tbl, ["MCSF Blount Island", "JEA Substation",              "MEDIUM", "Immediate",  "Storm surge, shoreline loss, water intrusion already active"])
table_row(tbl, ["NAS Jacksonville", "JEA Sub (north of base)",       "MEDIUM", "Immediate",  "Lack of redundancy; TECO gas as second SPOF"])
table_row(tbl, ["Camp Blanding",   "FPL Sub (north of base)",        "HIGH",   "Near-Term",  "Lack of redundancy, wildfire on transmission corridor"])
doc.add_paragraph()

note(
    "GIS would visualize these confirmed ratings spatially — overlaying substation locations "
    "against SLOSH surge zones, CDBG flood models, and USDA wildfire data. That analysis has "
    "not yet been performed by this team."
)

heading("1.3  Additional Confirmed Vulnerabilities from MIRR Documents", level=2)
confirmed = [
    ("Water intrusion already threatening electrical systems at Blount Island — ",
     "confirmed in the MCSF Blount Island Kickoff Briefing (April 2025). This is a current "
     "degradation, not a future scenario."),
    ("Camp Blanding depends on four separate power providers plus Seminole Electric — ",
     "confirmed in the Camp Blanding Summary V2. This creates a coordination complexity that "
     "is itself a vulnerability independent of physical hazard exposure."),
    ("JEA has initiated grid hardening and a microgrid feasibility study at NAS Jacksonville — ",
     "confirmed in the NAS JAX Summary V3. This is an existing effort, not a proposed action."),
    ("P035 Dual Fuel Generator Project at Blount Island is ERCIP-funded and in Final Design — ",
     "confirmed in the MCSF Blount Island Kickoff Briefing. Construction has not yet begun."),
]
for label, rest in confirmed:
    mixed(label, rest)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 2 — FRAMING
# ═════════════════════════════════════════════════════════════════════════════
heading("2. How We Are Framing This Problem", level=1)
body(
    "Based on reviewing the MIRR documents and the APTIM evaluation memo, the intern team "
    "identifies this as an energy risk management problem — not an energy efficiency problem. "
    "The installations are not wasting power. The supply chain delivering that power is fragile. "
    "Efficiency measures reduce consumption; they do not change the fact that a single substation "
    "failure cuts power to the entire installation."
)
body(
    "A useful frame drawn from the intern team's review: these installations currently function "
    "as isolated but dependent islands — drawing from the regional grid with no backup path. "
    "The APTIM memo (May 2026, Section 1.1) and the Hurricane Michael / Tyndall AFB experience "
    "both point toward the same long-term direction: installations need to become connected hubs "
    "capable of both drawing from and contributing to the regional grid. That framing shapes the "
    "investigation routes in Section 4."
)
note(
    "This framing is the intern team's analytical observation based on document review. "
    "It is not a MIRR-confirmed finding and has not been tested against GIS or operational data."
)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 3 — CASCADE MECHANISM
# ═════════════════════════════════════════════════════════════════════════════
heading("3. The Cascade: How Flooding Becomes an Energy Crisis", level=1)
body(
    "The following cascade is a logical analysis built on the confirmed MIRR findings above. "
    "It is not an independent finding — it connects dots that the MIRR documents already confirm "
    "individually. GIS analysis would test and quantify each link in this chain."
)

steps = [
    "Hazard event occurs — storm surge, compound rainfall flooding, or wildfire",
    "Transportation corridors flood or close:\n"
    "     A1A (NS Mayport)  |  Heckscher Drive (Blount Island)  |  "
    "Roosevelt Blvd / US-17 (NAS JAX)  |  SR-16 (Camp Blanding)\n"
    "     [All confirmed as transportation SPOFs in MIRR VA, Jan 2026 SC]",
    "Utility restoration crews cannot reach the affected substation",
    "Substation damage cannot be repaired — outage duration extends beyond "
    "what backup systems can cover",
    "Mission-critical systems lose power — mission disruption follows",
]
for i, s in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"Step {i}:  ")
    font(r1, bold=True)
    r2 = p.add_run(s)
    font(r2)

doc.add_paragraph()
body("Additional mechanisms that the MIRR documents suggest but that need further investigation:")

additional = [
    ("Saltwater corrosion: ",
     "Storm surge deposits salt residue on substation equipment even without full inundation, "
     "accelerating long-term degradation. Relevant to NS Mayport and Blount Island. "
     "Not yet spatially analyzed by this team."),
    ("Fuel delivery disruption: ",
     "Backup generator fuel is delivered by truck on the same corridors that block restoration "
     "crews. Duration of flooding relative to generator fuel capacity has not been assessed."),
    ("Water supply contamination: ",
     "Flooding in coastal and industrial areas can mobilize contaminants. The proximity of "
     "legacy industrial land uses near Blount Island and NAS JAX makes this worth investigating. "
     "No GIS overlay of contamination sites vs. flood zones has been done yet."),
    ("US-17 transmission corridor (unverified): ",
     "The intern team's prior research flagged transmission lines along US-17 as a common "
     "regional SPOF. This has not been verified through GIS or utility data. It is raised "
     "as a question, not a confirmed finding."),
]
for label, rest in additional:
    mixed(label, rest)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 4 — MISSION-CRITICAL CONTEXT
# ═════════════════════════════════════════════════════════════════════════════
heading("4. Why This Matters: Mission-Critical Context", level=1)
body(
    "The following is contextual background drawn from publicly available information about "
    "each installation's mission. It is included to help frame the severity of the cascade "
    "risk identified above — not as operational intelligence or classified findings."
)

installs = [
    ("NS Mayport — 4th Fleet Headquarters",
     "Home of the U.S. Navy's 4th Fleet, commanding operations across the Caribbean, Central "
     "America, and South America. Shore power to docked vessels, fueling operations, and command "
     "communications all depend on the JEA Mayport Substation — the highest-urgency confirmed "
     "SPOF in the region. The cascade risk here is direct: if the substation fails and A1A floods, "
     "those functions are on backup power with a finite fuel window."),
    ("NAS Jacksonville — Maritime Patrol Hub",
     "Primary East Coast hub for P-8 Poseidon maritime patrol aircraft. Aircraft maintenance, "
     "fueling, and pre-flight systems require continuous power. NAS JAX also has a second "
     "confirmed SPOF — the TECO privatized natural gas system — which is independent of the "
     "JEA electrical grid and not captured in the substation analysis."),
    ("MCSF Blount Island — Maritime Prepositioning Force",
     "Prepositioned Marine Corps equipment — vehicles, weapons systems, supplies — ready to "
     "deploy globally within 72 hours. Crane and conveyor operations, climate-controlled "
     "equipment storage, and logistics coordination all require continuous power. Water "
     "intrusion is already active at this site. A power outage here delays deployment "
     "capability — a consequence that scales with outage duration."),
    ("Camp Blanding — National Guard Mobilization Hub",
     "Florida's primary National Guard training and mobilization center. Four-provider "
     "coordination complexity is confirmed. A wildfire cutting the FPL transmission corridor "
     "could disrupt power without any weather warning, during a mobilization event, with "
     "limited pre-event preparation time."),
]
for title, desc in installs:
    heading(title, level=2)
    body(desc)

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 5 — INVESTIGATION ROUTES
# ═════════════════════════════════════════════════════════════════════════════
heading("5. Routes Worth Investigating", level=1)
body(
    "These are questions and avenues the intern team believes are worth pursuing — not "
    "proposed solutions. Each is grounded in something the MIRR documents raise or imply, "
    "but none has been verified through GIS or independent data collection yet."
)

heading("5.1  Spatial Verification of the Cascade Chain", level=2)
body("The cascade described in Section 3 needs GIS to move from logical to verified:")
cascade_qs = [
    "Do the confirmed transportation SPOFs (A1A, Heckscher Drive, Roosevelt Blvd, SR-16) "
    "actually intersect with the SLOSH Cat 3/5 or CDBG flood zones? By how much, and at "
    "what storm frequency?",
    "Are the confirmed substations within those same flood zones? What is their elevation "
    "relative to modeled surge and rainfall flood depths?",
    "Are the US-17 transmission lines a confirmed regional SPOF, or is this an assumption "
    "from the intern team's prior research that needs verification?",
]
for q in cascade_qs:
    bullet(q)
note("Datasets that would answer these: HIFLD substations, NHC SLOSH, NEFRC CDBG flood model, "
     "EIA transmission lines, FDOT road network. All are in the confirmed MIRR dataset inventory.")

heading("5.2  Tidal and River Current Energy — Is It Worth a Feasibility Study?", level=2)
body(
    "The St. Johns River mouth at NS Mayport has one of the highest tidal flow rates in Florida. "
    "Tidal energy converters sit below storm surge and generate on predictable cycles. This is "
    "not in the MIRR documents or the APTIM memo — it is an avenue the intern team identified "
    "as potentially relevant given the installation's location."
)
body("Questions that would need to be answered before this goes further:")
tidal_qs = [
    "Has the DOE Water Power Technologies Office or ESTCP already commissioned a feasibility "
    "study for tidal generation at Mayport or Blount Island?",
    "Do navigation channel requirements for the Mayport Ferry and military vessels rule out "
    "submerged turbines in the river mouth?",
    "What are the environmental permitting requirements for manatee and fish passage in "
    "the St. Johns River?",
]
for q in tidal_qs:
    bullet(q)
note("This cannot move beyond a question until a site-specific feasibility review is done. "
     "It is raised here because it is an option not mentioned in the APTIM memo and the "
     "location makes it worth asking about.")

heading("5.3  Landfill Methane as an Alternative Generation Source", level=2)
body(
    "The intern team's prior research flagged that MCAS Miramar (California) uses a blend "
    "of energy sources including landfill methane to support extended islanded operations. "
    "The Jacksonville region has multiple landfill sites in proximity to NAS Jacksonville "
    "and Blount Island."
)
body("Two things need to happen before this moves beyond a question:")
lf_qs = [
    "Source verification: the specific Miramar islanding duration and landfill methane "
    "details cited in the intern team's prior research have not been confirmed against a "
    "primary source. The APTIM memo only confirms that Miramar has 'microgrid and energy "
    "resilience investments' (Section 1.2, citing DOE 2022) — it does not mention landfill methane.",
    "GIS mapping: proximity of Jacksonville-area landfills to the installations, and whether "
    "pipeline routing would be feasible, has not been assessed.",
]
for q in lf_qs:
    bullet(q)
note("This is an honest area of uncertainty. It is worth verifying the source and mapping "
     "the geography before deciding whether to pursue it further.")

heading("5.4  Solar PV and Battery Storage at Camp Blanding", level=2)
body(
    "Camp Blanding's 73,000 acres and rural land availability raise the question of whether "
    "utility-scale solar under Enhanced Use Lease authority, paired with battery storage, "
    "could reduce dependence on the FPL transmission corridor. Solar PV and battery storage "
    "costs have declined significantly in recent years — noted in the intern team's prior "
    "research as a factor that may have changed the feasibility of options not previously "
    "considered viable."
)
body("What would need to be assessed:")
solar_qs = [
    "Has a solar feasibility study been commissioned at Camp Blanding under ESTCP or ESPC?",
    "What portion of Blanding's load could realistically be offset by on-site generation?",
    "Does the four-provider coordination structure complicate or simplify an EUL solar arrangement?",
]
for q in solar_qs:
    bullet(q)

heading("5.5  Civilian Infrastructure Dependency — How Deep Is the Gap?", level=2)
body(
    "All four installations depend on civilian-owned substations with no formal resilience "
    "obligation under current IGSAs. The MIRR has confirmed this gap. What is not yet known:"
)
dep_qs = [
    "What are the actual utility restoration timelines JEA and FPL commit to for comparable "
    "major outage events — and how do those compare to mission-critical power restoration "
    "requirements at each installation?",
    "Has either utility provided any data on planned substation hardening investments for "
    "the substations serving these installations?",
    "What would an expanded IGSA with resilience provisions actually need to include to be "
    "operationally meaningful?",
]
for q in dep_qs:
    bullet(q)
note("These are questions for utility coordination discussions and NAVFAC engagement — "
     "not answerable through GIS alone.")

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 6 — WHAT GIS WOULD ADD
# ═════════════════════════════════════════════════════════════════════════════
heading("6. What GIS Analysis Would Contribute", level=1)
body(
    "GIS does not create the vulnerabilities identified in this document — those are already "
    "confirmed by the MIRR process. What GIS adds is:"
)
gis_value = [
    "Spatial confirmation — showing exactly where confirmed SPOFs sit relative to confirmed "
    "hazard zones, moving from qualitative descriptions to measurable overlap",
    "Quantification — flood depth at substation locations, distance from SPOF to nearest "
    "alternative access route, extent of surge inundation along restoration corridors",
    "Visual communication — maps that make the cascade mechanism legible to decision-makers "
    "who may not engage with tabular vulnerability data",
    "New findings — spatial relationships the documents imply but don't confirm, such as "
    "the US-17 transmission corridor claim and the landfill proximity question",
]
for g in gis_value:
    bullet(g)

doc.add_paragraph()

tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = "Table Grid"
table_header(tbl2, ["Analysis", "Datasets Needed", "What It Would Confirm"])
gis_rows = [
    ("Substation vs. flood zone overlay",
     "HIFLD substations, NHC SLOSH, NEFRC CDBG",
     "Which substations are within Cat 3/5 surge or 100-yr flood extents"),
    ("Transportation SPOF vs. flood zone",
     "FDOT road network, same flood layers",
     "Which access corridors close at which storm frequencies"),
    ("US-17 transmission corridor",
     "EIA transmission lines, HIFLD",
     "Whether transmission lines along US-17 serve MIRR installations"),
    ("FPL corridor wildfire exposure",
     "USDA FS wildfire WUI, EIA transmission lines",
     "Extent of wildfire hazard along Camp Blanding's supply corridor"),
    ("Landfill proximity mapping",
     "EPA facility data, FDEP, installation boundaries",
     "Distance and routing feasibility for methane pipeline investigation"),
    ("Contamination sites vs. flood zones",
     "EPA ECHO, FDEP, FEMA NFHL",
     "Flood-mobilized contamination risk near Blount Island and NAS JAX"),
]
for r in gis_rows:
    table_row(tbl2, r)

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# SECTION 7 — SOURCES
# ═════════════════════════════════════════════════════════════════════════════
heading("7. Sources", level=1)
body("All vulnerability findings cited in this document are sourced to the following MIRR "
     "project materials. No independent GIS analysis has been performed by this team.")

sources = [
    "NEFRC MIRR Vulnerability Assessment — January 2026 Steering Committee",
    "MIRR Mutual Support Assessment — February 2026 TAC",
    "MIRR Adaptation Planning Framework — April 2026 Steering Committee",
    "MCSF Blount Island Kickoff Briefing — April 2025",
    "NAS Jacksonville Site Visit Summary V3",
    "Camp Blanding Summary V2",
    "NEFRC MIRR Stakeholder Workshop 1 Summary V2",
    "APTIM Memo: NE MIRR Evaluating Potential Adaptation Solutions — May 22, 2026",
    "GAO Climate Resilience Report — GAO-22-105452, 2021",
    "DoD Climate Adaptation Plan — 2021",
    "Prior Intern Research: Energy System Points of Failure and Geospatial Dataset Inventory",
]
for s in sources:
    bullet(s)

# ─────────────────────────────────────────────────────────────────────────────
out = "/home/user/HERMES/resilience_review/MIRR_Energy_Resilience_Contribution.docx"
doc.save(out)
print(f"Saved: {out}")
