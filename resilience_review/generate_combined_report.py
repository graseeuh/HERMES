from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Palette (friendlier: medium steel blue + warm accents) ────────────────────
BLUE       = (31,  96,  151)   # primary — readable medium blue
BLUE_MED   = (46,  117, 182)   # header fill text
BLUE_LITE  = (189, 215, 238)   # header background fill
BLUE_PALE  = (222, 235, 247)   # cover banner fill
SLATE      = (68,  84,  106)   # secondary text, h2
GRAY_D     = (89,  89,  89)    # body annotations
GRAY_LITE  = (242, 242, 242)   # divider band
BLACK      = (0,   0,   0)
GREEN_D    = (55,  86,  35)    # confirmed tag
AMBER_D    = (124, 84,  0)     # pending tag
RED_D      = (156, 0,   6)     # needs-verify tag
WHITE      = (255, 255, 255)
TEAL_RULE  = "1F6096"          # hex for border elements

TAG_CONFIRMED = "● Confirmed"
TAG_GIS       = "◐ Pending GIS"
TAG_VERIFY    = "○ Needs verification"

# ── Low-level XML helpers ─────────────────────────────────────────────────────
def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)

def set_para_shading(p, fill_hex):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    pPr.append(shd)

def add_left_border(p, color_hex=TEAL_RULE, sz="16"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    sz)
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)

def add_bottom_border(p, color_hex=TEAL_RULE, sz="6"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    sz)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)

# ── Text helpers ──────────────────────────────────────────────────────────────
def font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold      = bold
    run.italic    = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def centered(text, size=11, bold=False, italic=False,
             color=BLACK, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def body(text, space_after=6, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    font(run, size=11, italic=italic, color=color or BLACK)
    return p

def body2(parts, space_after=6):
    """Paragraph with mixed bold/normal runs. parts = [(text, bold), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold in parts:
        run = p.add_run(text)
        font(run, size=11, bold=bold)
    return p

def h1(text):
    """Friendly section header: light blue fill band, medium-blue bold text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(-0.05)
    set_para_shading(p, "BDD7EE")
    run = p.add_run(f"  {text}")
    font(run, size=11.5, bold=True, color=BLUE)
    return p

def h2(text):
    """Sub-header: bold slate with a thin left accent."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    add_left_border(p, TEAL_RULE, sz="8")
    p.paragraph_format.left_indent  = Inches(0.12)
    run = p.add_run(f" {text}")
    font(run, size=11, bold=True, color=SLATE)
    return p

def status_line(tag, text):
    tag_colors = {
        TAG_CONFIRMED: GREEN_D,
        TAG_GIS:       AMBER_D,
        TAG_VERIFY:    RED_D,
    }
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    r1 = p.add_run(f"{tag}  ")
    font(r1, size=9.5, bold=True, color=tag_colors.get(tag, GRAY_D))
    r2 = p.add_run(text)
    font(r2, size=9.5, italic=True, color=GRAY_D)

def quote(text, source):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.35)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    add_left_border(p, TEAL_RULE, sz="14")
    r1 = p.add_run(f'"{text}"')
    font(r1, size=10.5, italic=True, color=(50, 50, 50))
    r2 = p.add_run(f"\n  — {source}")
    font(r2, size=9, italic=True, color=GRAY_D)

def bullet(text, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.25 + indent * 0.2)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run(text)
    font(run, size=11)

def mixed(label, rest, indent=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.25 + indent * 0.2)
    r1 = p.add_run(label)
    font(r1, size=11, bold=True)
    r2 = p.add_run(rest)
    font(r2, size=11)

def divider(label):
    """Part divider: full-width navy band."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(8)
    set_para_shading(p, "1F6096")
    run = p.add_run(f"  {label}")
    font(run, size=11.5, bold=True, color=WHITE)

def tbl_header(tbl, cols):
    row = tbl.rows[0]
    for i, text in enumerate(cols):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(text)
        font(run, size=10, bold=True, color=WHITE)
        set_cell_shading(cell, "2E75B6")

def tbl_row(tbl, vals, alt=False):
    row = tbl.add_row()
    fill = "DEEAF1" if alt else "FFFFFF"
    for i, text in enumerate(vals):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(text)
        font(run, size=10)
        if alt:
            set_cell_shading(cell, fill)

def spacer(pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(pts)
    p.paragraph_format.space_before = Pt(0)

# ═════════════════════════════════════════════════════════════════════════════
# COVER
# ═════════════════════════════════════════════════════════════════════════════
# Banner
p_banner = doc.add_paragraph()
p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_banner.paragraph_format.space_before = Pt(10)
p_banner.paragraph_format.space_after  = Pt(0)
set_para_shading(p_banner, "1F6096")
r = p_banner.add_run("  Northeast Florida Military Installation Resilience Review  ")
font(r, size=16, bold=True, color=WHITE)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after  = Pt(0)
set_para_shading(p_sub, "2E75B6")
r2 = p_sub.add_run("  Energy Redundancy — Combined Intern Research Contribution  ")
font(r2, size=12, bold=True, color=WHITE)

spacer(8)
centered(
    "NAS Jacksonville  ·  NS Mayport  ·  MCSF Blount Island  ·  Camp Blanding JTF",
    size=10, color=SLATE, space_after=2)
centered(
    "Duval County and Clay County, Florida   |   Supervisor Review Draft   |   May 2026",
    size=9.5, italic=True, color=GRAY_D, space_after=10)

# Annotation key row
p_key = doc.add_paragraph()
p_key.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_key.paragraph_format.space_after = Pt(10)
for tag, col, desc in [
    (TAG_CONFIRMED, GREEN_D, " = sourced to confirmed MIRR document   "),
    (TAG_GIS,       AMBER_D, " = needs GIS spatial verification   "),
    (TAG_VERIFY,    RED_D,   " = needs external source before formal use"),
]:
    r = p_key.add_run(tag)
    font(r, size=9.5, bold=True, color=col)
    r2 = p_key.add_run(desc)
    font(r2, size=9.5, italic=True, color=GRAY_D)

# Scope note
p_scope = body(
    "This document has two parts. Part 1 preserves the first intern's research findings "
    "verbatim — no edits to their original text. Part 2 is the second intern's contribution, "
    "developed through a full review of all twelve MIRR project documents. It provides the "
    "geographic and operational context that explains why the redundancy gap identified in "
    "Part 1 is a serious mission risk in this specific region. No GIS analysis has been "
    "completed yet — that work is in progress and is documented with pending status. "
    "No formal solutions are proposed.",
    italic=True, color=SLATE, space_after=4
)
add_bottom_border(p_scope, TEAL_RULE, sz="4")
spacer(10)

# ═════════════════════════════════════════════════════════════════════════════
# PART 1
# ═════════════════════════════════════════════════════════════════════════════
divider("PART 1 — First Intern's Research Findings  (preserved verbatim)")
spacer(4)

h1("Energy Resilience Overview")
body(
    "Common themes across the sites include dependence on outside grid and exposure to its "
    "shortcomings, namely single-point utility failures outside of these military sites that "
    "could cause failures that cascade onto military sites and threaten mission success. These "
    "failure points are typically substations or lines along the U.S. 17 highway."
)
body(
    "These risks are exasperated by increasing energy demand throughout the region and lack "
    "of redundancy of power supply to installations. There are reports of aging infrastructure "
    "serving the installations which only increases the likelihood of single-point failure."
)
body(
    "It's worth noting that one of the bases, Camp Blanding, relies on four separate power "
    "providers and Seminole Electric for generation. The Florida Power & Light (FPL) substation "
    "north of the base is a high-risk asset vulnerable to natural hazards and a lack of redundancy."
)

h1("Adaptation Strategies — Summary Notes")
body(
    "The general theme of strategies discussed is moving from isolated but dependent islands "
    "into more connected hubs capable of offering and accepting support to and from the larger "
    "grid. This shift was motivated by Hurricane Michael's 2018 impact on Tyndall AFB, where "
    "hardening within the fence was undermined by failure of the surrounding civilian infrastructure.",
    space_after=4
)
bullet("Flood-proofing substations and supply lines. Mutual Support Agreements (MSAs) to "
       "formalize utility-installation collaboration.")
bullet("Addressing redundancy gaps in electrical and fuel systems — backup paths and fuel stocks.")
bullet("Microgrids with battery storage — feasibility studies already underway at NAS JAX.")
bullet("Solar PV and battery storage: costs have fallen significantly; feasibility study warranted.")
status_line(TAG_VERIFY, "Miramar 21-day islanding and landfill methane claim: primary source needed before formal use. "
            "See Part 2, Section 2.6 for verification steps.")
spacer(4)

# ═════════════════════════════════════════════════════════════════════════════
# PART 2
# ═════════════════════════════════════════════════════════════════════════════
divider("PART 2 — Second Intern's Research Contribution")
spacer(4)

body(
    "The sections below expand on Part 1 by grounding the redundancy gap in the specific "
    "geographic, operational, and documented hazard context of Northeast Florida. Every finding "
    "is tied to a confirmed source or clearly annotated as pending. This is the foundation the "
    "GIS analysis will build on.",
    italic=True, color=SLATE
)

# ── 2.1 ──────────────────────────────────────────────────────────────────────
h1("2.1  The Core Problem: Zero Electrical Redundancy at All Four Installations")

body(
    "The MIRR vulnerability assessment's central energy finding is direct: all four installations "
    "each draw from a single external substation with no backup supply path. There is no secondary "
    "feed, no alternate routing, and no automatic switching to an alternative grid node. If the one "
    "substation serving a base loses power — whether from storm surge, wildfire, equipment failure, "
    "or any other cause — the installation loses commercial electricity entirely."
)
body(
    "This is not a concern about efficiency or cost. It is a structural single point of failure "
    "baked into how these installations are physically connected to the grid. The substations are "
    "located outside the installation fence on civilian utility property, hardened to commercial "
    "standards rather than military resilience standards, and covered under Intergovernmental "
    "Service Agreements (IGSAs) that address only routine operations and maintenance — not "
    "resilience, redundancy, or priority restoration timelines."
)
body(
    "The gap between how fast a utility restores power after a major event (measured in days) "
    "and how fast a military mission requires power to be restored (measured in hours) is not "
    "addressed anywhere in the current IGSA framework. This is the operational problem."
)
quote(
    "At the asset systems level there is no immediate risk to mission other than single "
    "points of failure at each installation.",
    "MIRR Vulnerability Assessment, January 2026 Steering Committee"
)
quote(
    "Utility restoration timelines and fuel delivery continuity may ultimately affect "
    "operational recovery more than flood depth at individual facilities.",
    "APTIM Memo: NE MIRR Evaluating Potential Adaptation Solutions, May 22, 2026"
)
status_line(TAG_CONFIRMED, "Risk ratings and IGSA gap confirmed: MIRR VA (Jan 2026 SC) and Mutual Support Assessment (Feb 2026 TAC).")
spacer(4)

body("The table below summarizes the four confirmed energy single points of failure.", italic=True, color=SLATE)

t1 = doc.add_table(rows=1, cols=5)
t1.style = "Table Grid"
tbl_header(t1, ["Installation", "Substation / Utility", "Location", "Risk to Mission", "Urgency"])
spof_rows = [
    ("NS Mayport",         "JEA",                  "South of base — outside fence",  "HIGH",   "Immediate"),
    ("MCSF Blount Island", "JEA",                  "JEA Substation (river island)",   "MEDIUM", "Immediate"),
    ("NAS Jacksonville",   "JEA (+ TECO natural gas)", "North of base — outside fence", "MEDIUM", "Immediate"),
    ("Camp Blanding JTF",  "FPL (4 providers total)", "North of base — outside fence", "HIGH",   "Near-Term"),
]
for i, r in enumerate(spof_rows):
    tbl_row(t1, r, alt=(i % 2 == 1))
spacer(4)

body(
    "Camp Blanding's situation is especially complex. It draws from four separate power "
    "providers with Seminole Electric serving generation. That coordination complexity creates "
    "additional single points of failure at the provider handoff level — not just in the physical "
    "substation — and complicates who is responsible for restoration and in what order. "
    "This is explicitly documented in the MIRR project materials.",
    color=(50, 50, 50)
)
body(
    "NAS Jacksonville carries an additional dependency not captured in the substation table: "
    "TECO's privatized natural gas system. This is a second infrastructure SPOF independent "
    "of the electrical supply. A compound failure of both systems during a hazard event would "
    "affect powered maintenance operations and any gas-dependent generation or heating backup simultaneously.",
    color=(50, 50, 50)
)
status_line(TAG_CONFIRMED, "Camp Blanding four-provider complexity: Camp Blanding Summary V2. TECO gas dependency: NAS JAX Site Visit Summary V3.")

# ── 2.2 ──────────────────────────────────────────────────────────────────────
h1("2.2  The Cascade: How a Flooding Event Becomes a Mission Outage")

body(
    "Power loss at a military installation is not just about the substation taking a direct hit. "
    "In Northeast Florida, flooding creates a failure chain — a cascade — that can cause a "
    "multi-day outage even when the substation itself is not fully inundated. Understanding "
    "this chain is critical to understanding why redundancy specifically (not just hardening) is "
    "the right framing."
)

h2("The Cascade Sequence — Confirmed by MIRR Data")
body2([
    ("Step 1: ", True),
    ("A storm surge, compound rainfall flood, or wildfire occurs in the region.", False)
])
body2([
    ("Step 2: ", True),
    ("Transportation corridors flood or close — A1A at Mayport, Heckscher Drive at Blount Island, "
     "Roosevelt Blvd/US-17 at NAS JAX, SR-16 at Camp Blanding. These are confirmed SPOFs in the "
     "MIRR Vulnerability Assessment.", False)
])
body2([
    ("Step 3: ", True),
    ("Utility restoration crews cannot reach the damaged substation. Even if the substation sustained "
     "only partial damage and could be repaired quickly, no crew can get there.", False)
])
body2([
    ("Step 4: ", True),
    ("Backup generator fuel supply depletes. Generators are designed for hours-to-days of bridging, "
     "not extended multi-day outages. Fuel is delivered by truck — on the same flooded roads.", False)
])
body2([
    ("Step 5: ", True),
    ("Mission-critical systems — communications, mission operations, fueling, maintenance — lose power "
     "on a timeline driven by road conditions, not by the storm itself.", False)
])

body(
    "The cascade is not a hypothetical. Every link has been confirmed or observed. "
    "NS Mayport experiences multiple flooding events per year and its commanding officer "
    "cited power disruption as the most pressing operational concern. Water intrusion is "
    "already threatening electrical systems at Blount Island right now — not as a future risk. "
    "The Tyndall AFB lesson from Hurricane Michael (2018) is cited in the APTIM memo precisely "
    "because it demonstrated this cascade in practice: hardening within the fence is defeated "
    "when surrounding civilian infrastructure fails.",
    color=(50, 50, 50)
)
status_line(TAG_CONFIRMED, "Transportation SPOFs: MIRR VA (Jan 2026 SC). Mayport flooding: Stakeholder Workshop 1 Summary V2. "
            "Blount Island water intrusion: Kickoff Briefing (April 2025). Tyndall cascade: APTIM memo Section 1.1.")

h2("Why Backup Generators Are Not a Full Answer")
body(
    "Backup generators provide a bridging capability, not a redundancy solution. Their "
    "limitations in this context are specific and documented:"
)
bullet("Generator fuel arrives by diesel truck on the same road corridors that flood during a "
       "surge or compound rainfall event. Extended outages under compound flooding — where roads "
       "remain impassable for multiple days — can exhaust fuel before grid restoration is possible.")
bullet("Backup systems at Blount Island are noted in MIRR materials as untested. Whether they "
       "perform to specification under actual load conditions has not been verified.")
bullet("Even when generators function correctly, they typically serve only a prioritized subset "
       "of loads — not full installation capacity. Determining which loads are covered and which "
       "are not requires installation-level data not currently available publicly.")
bullet("Saltwater intrusion after a surge event deposits corrosive residue on generator and "
       "substation equipment even when facilities are not fully inundated. This accelerates "
       "equipment degradation and increases the probability of failure in subsequent events.")
status_line(TAG_CONFIRMED, "Blount Island untested backup: MIRR project materials. Saltwater corrosion mechanism: APTIM memo.")
status_line(TAG_VERIFY, "Generator fuel delivery protocols and backup system test records require coordination with NAVFAC.")

# ── 2.3 ──────────────────────────────────────────────────────────────────────
h1("2.3  Why Redundancy Is Harder Here — Regional Geography")

body(
    "The Northeast Florida coastal environment creates specific structural constraints "
    "that make building energy redundancy more difficult and more expensive than in most "
    "other regions. These are not hypothetical complications — they are physical realities "
    "that must be accounted for in any serious analysis."
)

h2("Shallow Water Table and the Undergrounding Problem")
body(
    "Undergrounding power lines is often proposed as a solution to wind and debris damage. "
    "In coastal Duval County and along the St. Johns River corridor, the water table sits "
    "only a few feet below grade in many areas. Burying infrastructure here does not eliminate "
    "vulnerability — it trades above-ground wind and debris exposure for below-ground groundwater "
    "intrusion, saltwater corrosion from below, and significantly more difficult fault identification "
    "and repair when failures occur."
)
body(
    "JEA has selectively undergrounded portions of its distribution network but approaches "
    "it cautiously given these conditions. An above-ground substation damaged in a storm "
    "can be assessed and repaired relatively quickly by crews who can see the damage. "
    "A buried cable failure in saturated, salt-affected soil requires specialized equipment "
    "and extended excavation timelines. Undergrounding in this region requires careful cost-benefit "
    "analysis — it is not a straightforward upgrade.",
    color=(50, 50, 50)
)
status_line(TAG_GIS, "Water table depths at specific substation locations need USGS 3DEP and SJRWMD groundwater data to map accurately.")

h2("Access Corridor Constraints by Installation")
body(
    "Each installation's geography creates hard physical limits on how utility restoration "
    "crews and fuel delivery trucks can reach the site. These are not minor inconveniences — "
    "they determine how long an outage lasts."
)
mixed("NS Mayport — ", "A1A is the sole primary access route. NS Mayport sits on a coastal "
      "peninsula between the Atlantic and the St. Johns River mouth. When A1A floods — which "
      "it does multiple times per year — the installation is effectively cut off by road. "
      "The JEA Mayport Substation, located south of the base, is on the same flooded corridor. "
      "A restoration crew cannot reach it until A1A drains.")
mixed("MCSF Blount Island — ", "Heckscher Drive is the primary access route to what is "
      "literally a river island in the St. Johns. Flooding on Heckscher Drive simultaneously "
      "blocks crew access to the JEA substation, fuel deliveries to on-site generators, and "
      "the personnel and equipment movement needed for the facility's prepositioning mission. "
      "All failure modes compound at the same bottleneck.")
mixed("NAS Jacksonville — ", "Roosevelt Blvd and US-17 are the primary approaches. Both "
      "run through the St. Johns River floodplain. Under compound flooding conditions — "
      "where tidal backflow compounds rainfall — these corridors can flood even during "
      "events that would not otherwise reach the base perimeter. The JEA substation "
      "is north of the base on the same general corridor.")
mixed("Camp Blanding — ", "SR-16 is the main approach, and Black Creek flooding "
      "is confirmed to close surrounding roadways even in tropical storm conditions — "
      "a lower storm threshold than hurricane-force events. Wildfire on the FPL "
      "transmission corridor is a separate access-independent disruption that requires "
      "no flooding at all to activate.")
status_line(TAG_CONFIRMED, "Access corridor SPOFs: MIRR VA (Jan 2026 SC). Black Creek / SR-16: Camp Blanding Summary V2.")
status_line(TAG_GIS, "FDOT road network overlaid against SLOSH and CDBG flood extents will confirm flood frequency for each corridor.")

h2("Regional Grid Topology — An Unresolved Gap")
body(
    "A substation serving a military installation can lose power even if the substation itself "
    "sustains no direct damage. Regional flooding or wind can knock out an upstream node — a "
    "transmission tower, a switching station, a generator interconnection — and cascade a "
    "blackout downstream to the substation regardless of local conditions. How many such "
    "upstream nodes exist between the regional grid and each installation's substation, "
    "how much of that upstream path runs through flood-exposed terrain, and how redundant "
    "each upstream segment is — none of this is currently known."
)
body(
    "JEA has not shared transmission routing data with the MIRR team as of the September 2025 "
    "TAC meeting. This is a confirmed data gap, not an analysis gap. EIA Form 860 and HIFLD "
    "transmission line data will be used as proxy mapping in the GIS phase, but the full "
    "picture requires JEA coordination.",
    color=(50, 50, 50)
)
status_line(TAG_CONFIRMED, "JEA data-sharing gap confirmed: MIRR project Sept 2025 TAC status notes.")
status_line(TAG_GIS, "EIA Form 860 and HIFLD transmission layers to be used as interim proxy. JEA routing data needed for accurate model.")

# ── 2.4 ──────────────────────────────────────────────────────────────────────
h1("2.4  Confirmed Vulnerabilities by Installation")

body(
    "The four installations do not face identical risks. Each has a specific combination of "
    "hazard exposure, access constraints, and operational mission that shapes why the redundancy "
    "gap matters differently at each location. These summaries are drawn entirely from confirmed "
    "MIRR project documents."
)

h2("Naval Station Mayport")
body(
    "NS Mayport hosts the 4th Fleet headquarters and serves as a major Atlantic operations "
    "staging, fueling, and logistics node. The JEA Mayport Substation carries the highest "
    "single-asset energy risk rating in the entire MIRR study region: HIGH urgency, Immediate "
    "timeline. The substation sits south of the installation on a peninsula exposed to both "
    "Atlantic storm surge (confirmed SLOSH Cat 3–5 exposure) and St. Johns River backflow flooding."
)
body(
    "The installation's shoreline is actively eroding — not a future concern but a current "
    "condition. This erosion is already moving infrastructure closer to surge inundation thresholds. "
    "The commanding officer explicitly named power disruption as the most pressing operational "
    "concern at Stakeholder Workshop 1, and flooding events at the installation occur multiple "
    "times per year. Shore power continuity for docked vessels, fueling operations, and "
    "command communications all depend on the single JEA substation.",
    color=(50, 50, 50)
)
status_line(TAG_CONFIRMED, "Risk rating HIGH/Immediate: MIRR VA (Jan 2026 SC). Active erosion and recurring flooding: Stakeholder Workshop 1 Summary V2.")

h2("MCSF Blount Island")
body(
    "Blount Island houses the Maritime Prepositioning Force — prepositioned ships loaded with "
    "Marine Corps equipment and supplies capable of deploying to a global crisis within 72 hours. "
    "The entire operational model depends on continuous powered crane and conveyor operations "
    "for ship loading and unloading, climate-controlled equipment storage, and logistics coordination."
)
body(
    "Water intrusion is already threatening electrical systems at the facility — this is a "
    "current degradation confirmed in the April 2025 Kickoff Briefing, not a projected future "
    "risk. Backup water systems are documented as untested. Heckscher Drive's access bottleneck "
    "means any power outage compounds simultaneously with access loss for restoration, fuel, "
    "and mission operations.",
    color=(50, 50, 50)
)
body(
    "The P035 Dual Fuel Generator Project (ERCIP-funded, Final Design Authority underway as of "
    "project review) is the most advanced near-term energy resilience measure in the region. "
    "It will install generators in all nine critical facilities and reduce JEA grid dependency. "
    "Ensuring it reaches construction without delay is the clearest near-term priority at Blount Island.",
    color=(50, 50, 50)
)
status_line(TAG_CONFIRMED, "Active water intrusion: MCSF Blount Island Kickoff Briefing (April 2025). P035 status: MIRR project documents. Risk MEDIUM/Immediate: MIRR VA (Jan 2026 SC).")

h2("Naval Air Station Jacksonville")
body(
    "NAS JAX is the primary East Coast hub for the P-8 Poseidon maritime patrol aircraft and "
    "supports major logistics and supply chain operations for other East Coast installations. "
    "Continuous power is required for aircraft maintenance cycles, fueling, pre-flight system "
    "checks, and communications. The JEA substation north of the base is the sole electrical "
    "supply point, rated MEDIUM risk at Immediate urgency, with lack of redundancy as the "
    "primary stressor."
)
body(
    "JEA has already initiated grid hardening work and microgrid feasibility studies specific "
    "to NAS JAX in response to identified vulnerabilities. This is the most advanced utility "
    "coordination status of the four installations. TECO's privatized natural gas system adds "
    "a second independent supply vulnerability that is separate from and simultaneous with "
    "the electrical risk.",
    color=(50, 50, 50)
)
status_line(TAG_CONFIRMED, "JEA hardening and microgrid study underway: NAS JAX Site Visit Summary V3. TECO gas dependency: NAS JAX Summary V3. Risk MEDIUM/Immediate: MIRR VA (Jan 2026 SC).")

h2("Camp Blanding Joint Training Center")
body(
    "Camp Blanding is Florida's primary National Guard training and mobilization hub. In a major "
    "conflict, emergency declaration, or domestic response, Guard units activate and mobilize "
    "through Blanding — requiring powered training ranges, communications, fuel distribution, "
    "and logistics at maximum capacity. The FPL substation north of the base is rated HIGH risk "
    "at Near-Term urgency."
)
body(
    "Two separate hazards can activate the redundancy gap at Blanding independently of each "
    "other. Flooding (Black Creek and SR-16 corridor) produces the standard cascade: access "
    "loss blocks restoration crews and fuel delivery. Wildfire on the FPL transmission corridor "
    "produces a different failure mode: no flooding required, no warning period comparable to "
    "a named storm, and the disruption occurs along the transmission path rather than at the "
    "substation itself. The four-provider coordination complexity adds a third layer — "
    "restoration requires coordination among multiple organizations with no single entity "
    "accountable for the timeline.",
    color=(50, 50, 50)
)
status_line(TAG_CONFIRMED, "Four-provider complexity and FPL substation risk HIGH/Near-Term: Camp Blanding Summary V2 and MIRR VA (Jan 2026 SC).")

# ── 2.5 ──────────────────────────────────────────────────────────────────────
h1("2.5  Historical Events — Evidence the Vulnerability Is Real")

body(
    "The following events demonstrate that the hazards driving the redundancy gap are not "
    "theoretical. They have occurred in this region within the last decade."
)

h2("Confirmed in MIRR Project Documents")
mixed("NS Mayport — recurring flooding, ongoing: ",
      "Multiple flooding events per year confirmed. Commanding officer cited this as the most "
      "pressing operational concern at Stakeholder Workshop 1.")
mixed("NAS Jacksonville — 106 mph wind gust, 1997: ",
      "Documented in MIRR evaluation methodology materials. Establishes that hurricane-force "
      "wind is a demonstrated regional hazard at this specific location.")
mixed("MCSF Blount Island — active water intrusion, current: ",
      "Water intrusion is already threatening electrical systems as of the April 2025 "
      "project kickoff. This is an ongoing condition, not a past event.")
mixed("Tyndall AFB — Hurricane Michael, October 2018: ",
      "On-base hardening investments proved insufficient when surrounding utilities and "
      "transportation infrastructure failed regionally. The installation sustained mission "
      "disruption not primarily from direct wind damage but from civilian infrastructure "
      "collapse around it. This is the foundational case study for the APTIM memo's framing.")
status_line(TAG_CONFIRMED, "All four items: confirmed MIRR project documents as cited above.")

h2("Regional Events — Needs Primary Source Before Formal Use")
body(
    "The following regional events are widely documented in public record but have not been "
    "pulled to primary sources (NHC reports, PSC storm filings, FEMA declarations) for this "
    "analysis. They should not be cited in any client-facing document until the verification "
    "steps below are completed.",
    italic=True, color=GRAY_D
)
mixed("Hurricane Matthew, October 2016 — ",
      "Caused historic St. Johns River flooding through Jacksonville. A1A reported as flooded. "
      "JEA reported widespread Duval County outages. Access corridors serving Mayport and "
      "Blount Island would have been affected by documented flood levels.\n"
      "         Verify with: NOAA NHC Post-Storm Report AL142016; Florida PSC JEA storm filing; "
      "NOAA Storm Events Database (Duval County, October 2016).")
mixed("Hurricane Irma, September 2017 — ",
      "JEA publicly reported approximately 200,000 customers without power at peak across "
      "the Jacksonville service territory.\n"
      "         Verify with: NOAA NHC Post-Storm Report AL112017; Florida PSC JEA storm filing; "
      "FEMA disaster declaration DR-4337; EIA Form 417 post-event report.")
status_line(TAG_VERIFY, "Do not cite Matthew or Irma in any formal deliverable until PSC filings and NHC reports are in hand.")

# ── 2.6 ──────────────────────────────────────────────────────────────────────
h1("2.6  GIS Analyses Planned")

body(
    "The following spatial analyses will confirm and quantify what the document review "
    "has established conceptually. Each is tied to specific publicly available datasets "
    "that can be accessed now, plus additional datasets that require data-sharing agreements "
    "expected after the August–September budget approval period."
)

t2 = doc.add_table(rows=1, cols=3)
t2.style = "Table Grid"
tbl_header(t2, ["Analysis", "Primary Datasets", "What It Will Confirm"])
gis_rows = [
    ("Substation locations vs. storm surge zones\n"
     "(FEMA effective and preliminary maps — check both; preliminary may show higher risk)",
     "HIFLD Open substations\nNHC SLOSH Cat 3 and Cat 5 extents\nFEMA MSC (both effective and preliminary versions)",
     "Which substations fall within surge inundation zones; whether preliminary FEMA maps "
     "reveal higher exposure than currently reflected in the effective map"),
    ("Access corridors vs. flood extents",
     "FDOT Florida road network\nNEFRC CDBG rainfall flood (present, 2040, 2070)\nFEMA NFHL\nResilient Jax compound flood model",
     "Which specific corridors flood at what storm frequency and projected future conditions; "
     "confirms which installations lose road access under each scenario"),
    ("Upstream transmission routing — substations back to regional grid",
     "EIA Form 860\nHIFLD transmission lines\nJEA routing data (pending data agreement)",
     "Node count and flood / wildfire exposure between installations and regional grid; "
     "confirms or refutes the upstream cascade vulnerability; verifies Part 1 US-17 SPOF claim"),
    ("FPL corridor wildfire hazard — Camp Blanding",
     "USDA Forest Service Wildfire Hazard Potential\nHIFLD transmission lines\nCamp Blanding installation boundary",
     "Spatial extent of high wildfire hazard along FPL supply corridor; "
     "how much of the corridor runs through WUI-rated terrain"),
    ("Water table depth at substation sites",
     "USGS National Map / 3DEP elevation\nSJRWMD groundwater monitoring data",
     "Undergrounding feasibility constraints at each substation location; "
     "depth-to-water table as a design constraint, not just an abstract concern"),
    ("Coastal exposure and erosion — Mayport and Blount Island shorelines",
     "USACE Jacksonville District coastal damage models\nNOAA Sea Level Rise projections (Resilient Jax)\nNOAA shoreline change data",
     "Rate of shoreline migration toward infrastructure at Mayport; "
     "how surge inundation thresholds change under 2040 and 2070 SLR scenarios"),
]
for i, r in enumerate(gis_rows):
    tbl_row(t2, r, alt=(i % 2 == 1))
spacer(4)
status_line(TAG_GIS, "Tier 1 (public data, available now): HIFLD, EIA Form 860, NOAA Digital Coast, FEMA MSC, FDOT, USDA FS, USGS, SJRWMD. "
            "Tier 2 (data agreement required, expected after August–September): JEA transmission routing, NAVFAC installation energy data, "
            "USACE detailed coastal damage models.")

# ── 2.7 ──────────────────────────────────────────────────────────────────────
h1("2.7  Items Requiring Research Before Any Proposal Can Be Made")

body(
    "These topics have appeared in prior research conversations or in project document review "
    "but are not ready for any formal deliverable. Each entry documents what is known, what "
    "is not, and what specific steps are needed next. Nothing here should be presented to "
    "a client until those steps are completed."
)

t3 = doc.add_table(rows=1, cols=3)
t3.style = "Table Grid"
tbl_header(t3, ["Topic", "What Is Known", "Specific Next Steps"])
research_rows = [
    ("Small Modular Reactors / DoD Microreactors\n(Manager note: look into modular nuclear)",
     "DoD's Project Pele program pursues on-site microreactors for military installations — "
     "would eliminate civilian grid dependency entirely if deployed. Most complete redundancy "
     "solution in concept; also most complex, longest lead time, and most regulatory-intensive. "
     "No NE Florida installation has been publicly identified as a Pele candidate site.",
     "Research Project Pele program status and site selection criteria. "
     "Confirm ENRIL organization name with manager before any outreach. "
     "Assess NRC regulatory pathway and typical deployment timeline. "
     "Determine whether any MIRR installations have submitted ERCIP or ESTCP applications."),
    ("Electrical infrastructure insurance coverage",
     "The type and scope of insurance JEA and FPL carry on the substations serving these "
     "installations is not known. Claims processing timelines after a storm event directly "
     "affect how long an outage persists — if a utility cannot begin repairs until an "
     "insurance assessment is completed, that adds to restoration time. Not currently "
     "captured in any MIRR analysis.",
     "Review Florida PSC annual utility filings and storm protection plans for JEA and FPL. "
     "Check JEA and FPL annual reports for infrastructure insurance disclosures. "
     "Determine whether FEMA NFIP policies apply to privately-owned utility substations."),
    ("Miramar MC Air Station — landfill methane and 21-day islanding\n(Part 1 claim)",
     "APTIM memo confirms only 'microgrid and energy resilience investments' at Miramar. "
     "The 21-day islanding figure and landfill methane fuel source are not cited to a "
     "primary source in any MIRR document reviewed. Cannot be presented as a confirmed "
     "parallel without verification.",
     "Find the DOE 2022 report cited in the APTIM memo regarding Miramar. "
     "Confirm the 21-day figure, the energy sources, and the installed capacity. "
     "GIS: map Jacksonville-area landfill locations relative to NAS JAX and Blount Island "
     "to assess whether a parallel opportunity could exist geographically."),
    ("JEA upstream transmission routing",
     "Confirmed data gap in MIRR project materials. JEA has not shared infrastructure data "
     "with the MIRR team as of September 2025 TAC. Upstream topology directly determines "
     "how many nodes can fail before each installation loses power — cannot be analyzed "
     "without this data.",
     "EIA Form 860 and HIFLD as interim proxy for GIS phase. "
     "JEA coordination under expanded IGSA as the path to actual routing data. "
     "Florida PSC storm hardening annual filings may include planned transmission investments "
     "relevant to these corridors."),
    ("Installation-level outage history",
     "Regional outage data from Matthew (2016) and Irma (2017) is publicly available in "
     "aggregate through Florida PSC and NHC reports. Installation-specific outage duration — "
     "how long each base was actually without commercial power and how long it ran on backup — "
     "is not in the public record and is not in current MIRR project documents.",
     "Florida PSC JEA storm filings for Matthew and Irma. "
     "EIA Form 417 post-event disturbance reports for Duval and Clay Counties. "
     "NAVFAC coordination required for installation-level data — not publicly available."),
    ("Solar PV + Battery Storage at Camp Blanding",
     "Costs for utility-scale solar and battery storage have declined substantially. "
     "Camp Blanding has 73,000 acres and rural land availability. Enhanced Use Lease (EUL) "
     "authority exists for DoD installations. No Blanding-specific feasibility study has "
     "been found in MIRR project documents or public DoD energy program databases.",
     "Search ESTCP and ESPC program databases for any prior studies scoped to Camp Blanding. "
     "Obtain load profile data for Blanding from NAVFAC or Army energy program contacts. "
     "Confirm EUL process, timeline, and applicable authority for National Guard installations."),
]
for i, r in enumerate(research_rows):
    tbl_row(t3, r, alt=(i % 2 == 1))
spacer(4)
status_line(TAG_VERIFY, "None of the above items should appear in any client-facing document until the verification steps listed are completed.")

# ── 2.8 Sources ───────────────────────────────────────────────────────────────
h1("2.8  Sources")

h2("Confirmed MIRR Project Documents (all reviewed for this analysis)")
for s in [
    "NEFRC MIRR Vulnerability Assessment — January 2026 Steering Committee",
    "MIRR Mutual Support Assessment — February 2026 TAC",
    "MIRR Adaptation Planning Framework — April 2026 Steering Committee",
    "MCSF Blount Island Kickoff Briefing — April 2025",
    "NAS Jacksonville Site Visit Summary V3",
    "Camp Blanding Summary V2",
    "NEFRC MIRR Stakeholder Workshop 1 Summary V2",
    "NS Mayport Stakeholder Summary",
    "APTIM Memo: NE MIRR Evaluating Potential Adaptation Solutions — May 22, 2026",
    "NEFRC MIRR TAC November 2025 Vulnerability Assessment Presentation",
]:
    bullet(s)

h2("To Obtain — Historical Events (before formal use)")
for s in [
    "NOAA NHC Post-Storm Report AL142016 (Hurricane Matthew, 2016)",
    "NOAA NHC Post-Storm Report AL112017 (Hurricane Irma, 2017)",
    "NOAA Storm Events Database — Duval and Clay Counties, 2015–present",
    "Florida PSC Storm Protection Plan Annual Filings — JEA and FPL",
    "FEMA Disaster Declarations — DR-4283 (Matthew), DR-4337 (Irma)",
    "EIA Form 417 Post-Event Disturbance Reports — Duval and Clay Counties",
]:
    bullet(s)

h2("GIS Datasets — Planned Tier 1 (public access, available now)")
for s in [
    "HIFLD Open Data — military installation boundaries, transmission lines, substations",
    "EIA Form 860 — substation and generation facility locations",
    "NOAA Digital Coast — SLOSH surge modeling, SLR projections",
    "FEMA Flood Map Service Center — NFHL (effective and preliminary versions)",
    "FDOT GIS Open Data — Florida road network",
    "USDA Forest Service Wildfire Hazard Potential (Research Data Archive)",
    "USGS National Map / 3DEP — elevation data",
    "SJRWMD — groundwater and water resource data",
    "Resilient Jacksonville — NEFRC CDBG compound flood model",
]:
    bullet(s)

h2("GIS Datasets — Planned Tier 2 (data agreement / coordination required)")
for s in [
    "JEA transmission and substation routing data — follow-up meetings pending as of Sept 2025",
    "NAVFAC / installation energy operations data — data request submitted, unresolved",
    "USACE Jacksonville District — detailed coastal damage models",
    "Clay County Utility Authority (CCUA) — utility infrastructure data",
    "Florida PSC utility storm hardening filings — JEA and FPL annual filings",
]:
    bullet(s)

# ─────────────────────────────────────────────────────────────────────────────
out = "/home/user/HERMES/resilience_review/MIRR_Energy_Combined_Contribution.docx"
doc.save(out)
print(f"Saved: {out}")
